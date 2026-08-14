"""Build LINK-screenshot-report.html + .docx from png/ + captions."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parent
PNG = ROOT / "png"
HTML_OUT = ROOT / "LINK-screenshot-report.html"
DOCX_OUT = ROOT / "LINK-screenshot-report.docx"
README_OUT = ROOT / "README.txt"

SHOTS = [
    # Consumer
    ("00-cookie-banner.png", "Consumer site", "Cookie / consent banner",
     "First-visit consent gate — documents LINK’s privacy posture for the course report."),
    ("01-home-hero.png", "Consumer site", "Home hero",
     "Brand-led first viewport with Family Pool iPhone — the core “worth more together” promise."),
    ("02-home-concept-mid.png", "Consumer site", "Home · concept mid-page",
     "Connect → Pool → Use storytelling that sells why scattered loyalty balances need LINK."),
    ("03-home-partners-cta.png", "Consumer site", "Home · partners / CTA band",
     "Lower home band with partner ecosystem and conversion CTAs."),
    ("04-how-it-works-piggy.png", "Consumer site", "How it works · piggy hero",
     "Piggy-bank hero for Connect. Pool. Use together — memorable report visual."),
    ("05-for-families.png", "Consumer site", "For families",
     "Household roles (Nan, Wit, Ploy) — who earns, who pools, who burns."),
    ("06-partners.png", "Consumer site", "Partners",
     "Partner value page for Lotus’s, AIS, BTS, iBerry, IHG — ecosystem proof."),
    ("07-earn-to-burn.png", "Consumer site", "Earn to burn",
     "Earn → pool → goal → burn loop linking everyday spend to a family getaway."),
    ("08-demo.png", "Consumer site", "Demo",
     "Interactive product demo surface for evaluators and partners."),
    ("09-home-hero-mobile.png", "Consumer site", "Home hero · mobile (390px)",
     "Phone-first Family Pool presentation for mobile report frames."),
    ("10-how-it-works-mobile.png", "Consumer site", "How it works · mobile (390px)",
     "Mobile piggy hero — responsive storytelling for the report."),
    # CDP
    ("11-cdp-login.png", "CDP Admin", "CDP Admin login",
     "Professor gate before the simulated CDP workspace (demo credentials)."),
    ("12-cdp-overview-kpis.png", "CDP Admin", "CDP overview · KPI strip",
     "Pilot KPIs with course rail — Unify → Segment → Activate → Measure."),
    ("13-cdp-identity-kent.png", "CDP Admin", "Identity graph · Kent OldKiwi",
     "Golden-record profile with photo, traits, and resolved partner IDs."),
    ("14-cdp-partner-tiles.png", "CDP Admin", "Partner mock data tiles",
     "Lotus’s, AIS, BTS, iBerry, IHG attribute tiles on Kent’s profile."),
    ("15-cdp-flag-for-ads.png", "CDP Admin", "Flag for ads · activation",
     "Activation moment: queue Kent to LINE / Meta / IHG offer destinations."),
    ("16-cdp-ai-propensity.png", "CDP Admin", "AI propensity · Q4 2026 travel",
     "Travel propensity panel — predictive segment that drives partner offers."),
    ("17-cdp-funnel.png", "CDP Admin", "Conversion funnel",
     "Join → Earn → Pool → Goal → Burn measurement for the earn-to-burn loop."),
    ("18-cdp-course-rail.png", "CDP Admin", "Course materials rail",
     "Sidebar course rail — academic framing alongside the live CDP demo."),
]


def compress_pngs(max_side: int = 1600) -> None:
    """Downscale and optimize PNGs for reasonable file size."""
    for path in sorted(PNG.glob("*.png")):
        if path.name.startswith("_"):
            continue
        try:
            im = Image.open(path)
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGBA")
        except Exception as e:
            print("skip", path.name, e)
            continue
        w, h = im.size
        if max(w, h) > max_side:
            im.thumbnail((max_side, max_side), Image.Resampling.LANCZOS)
        before = path.stat().st_size
        # Prefer JPEG for photographic UI shots when savings are large
        if before > 700_000:
            rgb = im.convert("RGB")
            jpg = path.with_suffix(".jpg")
            rgb.save(jpg, format="JPEG", quality=82, optimize=True)
            if jpg.stat().st_size < before * 0.7:
                path.unlink()
                print(f"{path.name} -> {jpg.name} ({before//1024}->{jpg.stat().st_size//1024}KB)")
                continue
            jpg.unlink(missing_ok=True)
        tmp = path.with_suffix(".tmp.png")
        im.save(tmp, format="PNG", optimize=True)
        after = tmp.stat().st_size
        if after < before:
            tmp.replace(path)
            print(f"opt {path.name} ({before//1024}->{after//1024}KB)")
        else:
            tmp.unlink(missing_ok=True)


def resolve_file(stem_png: str) -> Path | None:
    p = PNG / stem_png
    if p.exists():
        return p
    j = PNG / stem_png.replace(".png", ".jpg")
    if j.exists():
        return j
    return None


def build_html(shots: list[tuple]) -> None:
    sections: dict[str, list] = {}
    for item in shots:
        sections.setdefault(item[1], []).append(item)

    parts = [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head>",
        '<meta charset="UTF-8" />',
        '<meta name="viewport" content="width=device-width, initial-scale=1" />',
        "<title>LINK screenshot report — marketing site &amp; CDP Admin</title>",
        "<style>",
        """
:root { --ink:#1a1a1a; --muted:#5a5a5a; --line:#e4e4e4; --bg:#fafafa; --accent:#2360e8; }
* { box-sizing: border-box; }
body { margin:0; font-family: "Segoe UI", system-ui, sans-serif; color:var(--ink); background:var(--bg); line-height:1.45; }
header.page { padding:2.2rem 1.5rem 1.2rem; max-width:980px; margin:0 auto; }
header.page h1 { font-size:1.75rem; margin:0 0 .4rem; }
header.page p { margin:.25rem 0; color:var(--muted); max-width:62ch; }
.note { font-size:.92rem; background:#fff; border:1px solid var(--line); padding:.85rem 1rem; border-radius:8px; margin-top:1rem; }
main { max-width:980px; margin:0 auto; padding:0 1.5rem 3rem; }
h2 { margin:2rem 0 .75rem; font-size:1.25rem; border-bottom:2px solid var(--accent); padding-bottom:.35rem; display:inline-block; }
figure { margin:0 0 1.75rem; background:#fff; border:1px solid var(--line); border-radius:10px; overflow:hidden; }
figure img { display:block; width:100%; height:auto; }
figcaption { padding:.75rem 1rem 1rem; }
figcaption strong { display:block; margin-bottom:.2rem; }
figcaption span { color:var(--muted); font-size:.95rem; }
@media print {
  body { background:#fff; }
  figure { break-inside: avoid; border:none; }
  header.page, .note { break-after: avoid; }
}
""",
        "</style>",
        "</head>",
        "<body>",
        '<header class="page">',
        "<h1>LINK screenshot report</h1>",
        "<p>Marketing site + CDP Admin frames for professor / LINK course reporting. Captured from "
        "<a href=\"https://link.afolio.co/\">link.afolio.co</a> "
        "(AI propensity &amp; partner attribute tiles from matching local build when live was behind).</p>",
        '<div class="note"><strong>Illustrative only.</strong> CDP demo login: <code>kent</code> / <code>2026</code>. '
        "Simulated pilot data — not production telemetry. Open this HTML in a browser; images use relative paths "
        "(<code>png/</code>) so you can copy-paste into Word or Google Docs.</div>",
        "</header>",
        "<main>",
    ]

    for section, items in sections.items():
        parts.append(f"<h2>{section}</h2>")
        for file, _sec, title, caption in items:
            resolved = resolve_file(file)
            if not resolved:
                continue
            rel = f"png/{resolved.name}"
            parts.append("<figure>")
            parts.append(f'<img src="{rel}" alt="{title}" loading="lazy" />')
            parts.append("<figcaption>")
            parts.append(f"<strong>{title}</strong>")
            parts.append(f"<span>{caption}</span>")
            parts.append("</figcaption>")
            parts.append("</figure>")

    parts.extend(["</main>", "</body>", "</html>"])
    HTML_OUT.write_text("\n".join(parts), encoding="utf-8")
    print("wrote", HTML_OUT)


def build_docx(shots: list[tuple]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("LINK screenshot report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run(
        "Marketing site + CDP Admin frames for professor / LINK course reporting. "
        "Source: https://link.afolio.co/ (selected CDP panels from matching local build when live lagged)."
    )

    note = doc.add_paragraph()
    run = note.add_run(
        "Illustrative only. CDP demo login: kent / 2026. Simulated pilot data — not production telemetry."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    current = None
    for file, section, heading, caption in shots:
        resolved = resolve_file(file)
        if not resolved:
            continue
        if section != current:
            current = section
            doc.add_heading(section, level=1)
        doc.add_heading(heading, level=2)
        # Width ~6.5" for readability in Word
        doc.add_picture(str(resolved), width=Inches(6.3))
        cap = doc.add_paragraph(caption)
        cap.runs[0].font.size = Pt(10)
        cap.runs[0].font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)

    doc.save(DOCX_OUT)
    print("wrote", DOCX_OUT)


def write_readme() -> None:
    README_OUT.write_text(
        """LINK report screenshot pack
============================

Contents
- png/  (or jpg siblings) — numbered report frames
- LINK-screenshot-report.html — open in a browser; copy images into Word/Docs
- LINK-screenshot-report.docx — same structure with embedded images

CDP Admin demo credentials (course / professor access)
  username: kent
  password: 2026

Notes
- Shots are illustrative / simulated pilot UI — not live production telemetry.
- Most consumer + CDP frames captured from https://link.afolio.co/
- AI propensity (16) and partner attribute tiles (14) used the matching local
  website build when those panels were not yet on the live deploy.
- Re-capture: node capture.mjs (from this folder; requires Playwright)

""",
        encoding="utf-8",
    )
    print("wrote", README_OUT)


def main() -> None:
    # Remove smoke test
    for junk in PNG.glob("_*.png"):
        junk.unlink()
        print("removed", junk.name)

    compress_pngs()
    available = []
    for item in SHOTS:
        if resolve_file(item[0]):
            available.append(item)
        else:
            print("MISSING", item[0])

    # Rewrite manifest cleanly
    (ROOT / "manifest.json").write_text(
        json.dumps(
            {
                "shots": [
                    {"file": resolve_file(f).name, "section": s, "title": t, "caption": c}
                    for f, s, t, c in available
                ]
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    build_html(available)
    build_docx(available)
    write_readme()

    total = sum(p.stat().st_size for p in PNG.iterdir() if p.is_file())
    print(f"image bytes: {total/1e6:.2f} MB")


if __name__ == "__main__":
    main()
